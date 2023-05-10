from resume_parser import resumeparse
from tika import parser
import locationtagger
import re
from nltk import ne_chunk, pos_tag, word_tokenize
from nltk.tree import Tree
from pdfminer.high_level import extract_text
import aspose.words as aw
import pandas as pd
import docx2txt
from PyPDF2 import PdfReader

import glob
from pdfminer.high_level import extract_text
import spacy
from spacy import displacy
nlp = spacy.load("en_core_web_lg")
import docx
import os
import pathlib
##############################################
base_dir = pathlib.Path(__name__).parent.absolute()
os.path.join(base_dir,'data_csv', 'client22.csv')
##############################################
def doc_convert_to_pdf(path):
    #     count =0
    cunt =0
    cnt = 0
    for file_ in glob.glob(path):
        
        try:
            if file_.endswith(".doc"):
                doc = aw.Document(file_)
                doc.save(os.path.join(base_dir,'pdf_file', f'{cunt}.pdf'))
                cunt += 1
            elif file_.endswith(".rtf"):
                doc = aw.Document(file_)
                doc.save(os.path.join(base_dir,'pdf_file', f'{cnt}.pdf'))
                cnt += 1
        except  Exception as e:
            print(e)


# convert word to text
def extract_text_from_word(word_path):
    doc = docx.Document(word_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# convert pdf to text
def extract_text_from_pdf(pdf_path):
    return extract_text(pdf_path)


# return text
def return_text(path):
    if path.endswith('.docx'):
        return extract_text_from_word(path)
    elif path.endswith('.pdf'):
        return extract_text_from_pdf(path)


# convert_the_info_dic
def convert_the_info_dic(text):
    # breaking the text
    dit = {}
    title = ['Education', 'Qualification', 'Skill', 'Objective', 'Experience', 'Tools', 'Designation', 'Employment']
    for i in title:
        text = text.replace(f'{i}', f'|{i}').replace(f'{i.upper()}', f'|{i.upper()}')
    text = text.split("|")

    # making dictionary of inforamtion
    for titl in title:
        check = False
        for txt in text:
            if titl in txt or titl.upper() in txt:
                check = True
                dit[titl] = txt
        if check == False:
            dit[titl] = 'Nan'
    return dit


# extraction_the_data
def extraction_of_Data(resume_path, data_in_dic, text):
    prev = []
    person = []
    city = []
    education = []
    doc = nlp(text)

    #     #extraction of person
    #     for ent in doc.ents:
    #         if ent.label_ == "PERSON" or (ent.text.istitle()):
    #             person.append(ent.text)
    #             break
    #     #extraction of city
    #     for ent in doc.ents:
    #         if ent.label_ == 'GPE' or  ent.label_ =='LOC':
    #             city.append(ent.text)
    #             break

    #     extraction of Education
    edu_attribute = data_in_dic['Education']
    education.append(edu_attribute)

    data1 = resumeparse.read_file(resume_path)
    # extract the name
    name_data = data1['name']
    if 'Aspose Pty' == name_data:
        people = []
        for ent in doc.ents:
            if ent.label_ == "PERSON" or (ent.text.istitle()):
                people.append(ent.text)
        name_data = people[2]

    else:
        name_data = name_data
    # extract the email
    email_data = data1['email']
    # extract the phone
    phone = data1['phone']
    if str(phone) == '003-2023':
        phone_pattern = re.compile(r'(?:\+\d{1,2}\s)?\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}')
        phone_numbers = re.findall(phone_pattern, text)
        phone = phone_numbers[0]
    else:
        phone = phone
    # extract the skills
    skill_data = data1['skills']
    # extract the designition
    designation_data = data1['designition']
    # extract the experience
    experience_data = data1['total_exp']
    if len([i for i in str(experience_data)]) > 2:
        experience_data = None
    else:
        experience_data = experience_data

    # append the information
    #     prev.append(person)
    prev.append(name_data)
    prev.append(email_data)
    prev.append(phone)
    prev.append(str(skill_data))
    prev.append(designation_data)
    prev.append(experience_data)
    prev.append(education)
    #     prev.append(city)
    # location
    place_entity = locationtagger.find_locations(text=text)
    prev.append(str(place_entity.countries))
    prev.append(str(place_entity.regions))
    prev.append(str(place_entity.cities))
    prev.append(resume_path)
    return prev



